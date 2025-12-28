import csv
import re
from docx import Document
import pywer


def normalize(text: str) -> str:
    # Remove leading line numbers and dots at each line start
    text = re.sub(r'^\s*\d+\.\s*', '', text, flags=re.MULTILINE)
    # Remove all digits (numbers)
    text = re.sub(r'\d+', '', text)
    # Remove all brackets and their angle variants
    text = re.sub(r'[\[\]\(\)\{\}\<\>]', '', text)
    # Convert to lowercase
    text = text.lower()
    # Collapse multiple whitespace characters into a single space
    text = re.sub(r'\s+', ' ', text)
    # Strip leading and trailing whitespace
    return text.strip()


# Load documents
hand_doc = Document("mestrija_handtranscription.docx")
trans_doc = Document("mestrija_transkribus.docx")


# Function to split document text by page breaks
def split_by_page(doc):
    pages = []
    current_page_text = []
    for para in doc.paragraphs:
        if para.text.strip() == '' and current_page_text:
            pages.append('\n'.join(current_page_text))
            current_page_text = []
        else:
            current_page_text.append(para.text)
    if current_page_text:
        pages.append('\n'.join(current_page_text))
    return pages


from docx import Document

hand_pages = split_by_page(hand_doc)
trans_pages = split_by_page(trans_doc)

def split_by_page(doc):
    pages = []
    current_page = []
    for para in doc.paragraphs:
        current_page.append(para.text)
        # Provjera da li odlomak sadrži prekid stranice
        if para._element.xpath('.//w:br[@w:type="page"]'):
            pages.append('\n'.join(current_page))
            current_page = []
    if current_page:
        pages.append('\n'.join(current_page))
    return pages

num_pages = min(len(hand_pages), len(trans_pages))


# Calculate CER and WER for each page
results = []
for i in range(num_pages):
    ref = normalize(hand_pages[i])
    hyp = normalize(trans_pages[i])
    cer = pywer.cer([ref], [hyp])
    wer = pywer.wer([ref], [hyp])
    results.append([i+1, cer, wer])


# Write results in CSV file
with open('page_level_cer_wer.csv', 'w', newline='', encoding='utf-8') as csvfile:
    writer = csv.writer(csvfile)
    writer.writerow(['Page', 'CER', 'WER'])
    writer.writerows(results)

print("Page-level CER and WER written to page_level_cer_wer.csv")

# Export .txt files
with open("reference_normalized.txt", "w", encoding="utf-8") as f:
    f.write(ref)
with open("hypothesis_normalized.txt", "w", encoding="utf-8") as f:
    f.write(hyp)
